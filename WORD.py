from docx import Document

# 创建Word文档
doc = Document()

# 添加标题
doc.add_heading('D-H密钥交换协议计算过程', 0)

# 添加计算步骤和结果
doc.add_heading('1. 选择素数和原根', level=1)
doc.add_paragraph(f"选择素数 p = 223, 原根 g = 5")

doc.add_heading('2. 私钥选择', level=1)
doc.add_paragraph(f"成员A选择私钥 a = {11}")
doc.add_paragraph(f"成员B选择私钥 b = {139}")

doc.add_heading('3. 公钥计算', level=1)
doc.add_paragraph(f"成员A计算公钥 A = g^a mod p = {45}")
doc.add_paragraph(f"成员B计算公钥 B = g^b mod p = {176}")

doc.add_heading('4. 共享密钥计算', level=1)
doc.add_paragraph(f"成员A计算共享密钥 K_A = B^a mod p = {145}")
doc.add_paragraph(f"成员B计算共享密钥 K_B = A^b mod p = {145}")

doc.add_heading('5. 结果验证', level=1)
doc.add_paragraph(f"共享密钥一致: {145}")

# 保存文档
doc.save("DH_密钥交换计算过程.docx")

print("计算过程已保存到 DH_密钥交换计算过程.docx")
